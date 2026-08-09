from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.config import ARTIFACTS_DIR, EXPERIMENTS_DIR, FIGURES_DIR, REPORTS_DIR

REPORT_PATH = REPORTS_DIR / "Informe_Proyecto_1_MLP.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
MUTED = RGBColor(90, 98, 108)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def font_run(run, size=11, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    font_run(run, size=9, color="666666")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = False
    font_run(p.add_run(text), size=9, italic=True, color="555555")


def add_figure(doc, filename, caption, width=6.05):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(FIGURES_DIR / filename), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        font_run(p.add_run(str(text)), size=font_size, bold=True)
    header_props = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_props.append(repeat)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i < 3 else WD_ALIGN_PARAGRAPH.RIGHT
            font_run(p.add_run(str(value)), size=font_size)
        row_props = table.rows[-1]._tr.get_or_add_trPr()
        row_props.append(OxmlElement("w:cantSplit"))
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.5)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    p.add_run(text)
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "CC3092 | Proyecto 1 - Competencia de Modelación"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    font_run(header.runs[0], size=9, color="6A7078")
    add_page_number(section.footer.paragraphs[0])


def build_report():
    eda = json.loads((ARTIFACTS_DIR / "eda_summary.json").read_text())
    best = json.loads((ARTIFACTS_DIR / "best_configuration.json").read_text())
    holdout = json.loads((ARTIFACTS_DIR / "holdout_metrics.json").read_text())
    missing = pd.read_csv(ARTIFACTS_DIR / "missing_summary.csv", index_col=0)
    corr = pd.read_csv(ARTIFACTS_DIR / "target_correlations.csv", index_col=0).iloc[:, 0]
    segments = pd.read_csv(ARTIFACTS_DIR / "error_by_price_segment.csv")
    largest = pd.read_csv(ARTIFACTS_DIR / "largest_errors.csv")
    results = pd.read_csv(EXPERIMENTS_DIR / "results.csv")
    trials = pd.read_csv(EXPERIMENTS_DIR / "optuna_trials.csv")
    confirm = best["confirmatory_cv"]
    params = best["params"]
    baseline = float(results.loc[results.experiment_id == "baseline_raw", "cv_rmse_mean"].iloc[0])
    improvement = (baseline - confirm["cv_rmse_mean"]) / baseline * 100

    doc = Document()
    configure_document(doc)

    # Editorial cover pattern, implemented as a restrained academic report cover.
    doc.add_paragraph().paragraph_format.space_after = Pt(92)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(kicker.add_run("CC3092 DEEP LEARNING Y SISTEMAS INTELIGENTES"), size=11, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(8)
    font_run(title.add_run("Competencia de modelación\ncon Multi-Layer Perceptron"), size=27, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    font_run(subtitle.add_run("Predicción de SalePrice en Ames Housing"), size=15, color="4D5A66")
    metric = doc.add_paragraph()
    metric.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(metric.add_run(f"RMSE CV: {confirm['cv_rmse_mean']:,.2f} +/- {confirm['cv_rmse_std']:,.2f} USD"),
             size=13, bold=True, color=BLUE)
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(90)
    font_run(author.add_run("Javier Valladares"), size=12, bold=True)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(date.add_run("Agosto de 2026"), size=11, color="666666")
    doc.add_page_break()

    doc.add_heading("Resumen ejecutivo", level=1)
    doc.add_paragraph(
        f"Se desarrolló un MLP para regresión sobre {eda['rows']:,} viviendas y {eda['columns'] - 1} predictores. "
        f"La configuración se seleccionó sin consultar el holdout, mediante exploración, {len(trials)} trials de Optuna "
        f"({int((trials.state == 'PRUNED').sum())} podados) y una confirmación estratificada de cinco folds. El modelo "
        f"ganador redujo el RMSE CV de {baseline:,.2f} a {confirm['cv_rmse_mean']:,.2f} USD ({improvement:.1f}%); "
        f"en el holdout interno independiente obtuvo {holdout['rmse']:,.2f} USD. El artefacto de competencia se "
        "reentrenó después con todas las observaciones disponibles. Todas las métricas se calcularon en la escala original de SalePrice."
    )
    doc.add_heading("Criterio de lectura de las métricas", level=2)
    doc.add_paragraph(
        "El RMSE penaliza de forma cuadrática los errores grandes. Por ello, una diferencia entre CV y holdout puede "
        "aparecer cuando cambia la concentración de viviendas de precio alto. El RMSE del artefacto final no se estima "
        "sobre las mismas filas usadas en su reentrenamiento; la cifra honesta procede del modelo entrenado solo con el 80% de desarrollo."
    )

    doc.add_heading("1. Análisis exploratorio de datos (EDA)", level=1)
    doc.add_heading("1.1 Dimensiones, tipos y objetivo", level=2)
    doc.add_paragraph(
        f"El CSV contiene {eda['rows']:,} observaciones y {eda['columns']} columnas: Id, {eda['numeric_features']} predictores "
        f"numericos, {eda['categorical_features']} categóricos y SalePrice. No hay filas ni Id duplicados. MSSubClass "
        "(15 códigos de tipo constructivo) y MoSold (12 meses) son números por almacenamiento, pero semánticamente son categorías; "
        "OverallQual y OverallCond se conservaron como escalas ordinales numéricas."
    )
    target = eda["target"]
    target_rows = [
        ("Media", f"{target['mean']:,.2f}"), ("Mediana", f"{target['median']:,.2f}"),
        ("Desviación estándar", f"{target['std']:,.2f}"), ("Mínimo", f"{target['min']:,.0f}"),
        ("Máximo", f"{target['max']:,.0f}"), ("Asimetría", f"{target['skew']:.3f}"),
    ]
    add_table(doc, ["Estadístico", "SalePrice (USD)"], target_rows, [3600, 5760], font_size=10)
    add_figure(doc, "target_distribution.png", "Figura 1. SalePrice original y transformado con log1p.")
    doc.add_paragraph(
        "La media supera a la mediana y el máximo alcanza 745,000 USD, evidencia de una cola derecha pronunciada. "
        "Esto justificó probar log1p, pero la transformación final se decidió por RMSE original, no por simetría visual."
    )

    doc.add_heading("1.2 Valores faltantes, atípicos e inconsistencias", level=2)
    top_missing = [(idx, int(row.missing_count), f"{row.missing_pct:.1f}%") for idx, row in missing.head(10).iterrows()]
    add_table(doc, ["Variable", "Faltantes", "%"], top_missing, [4500, 2460, 2400], font_size=9.5)
    doc.add_paragraph(
        f"Existen {eda['missing_cells']:,} celdas faltantes repartidas en {eda['columns_with_missing']} columnas. En variables como "
        "PoolQC, Alley, Fence, chimenea, garage y sótano, la ausencia suele representar que el elemento no existe, por lo que "
        "se preservó como categoría explícita Missing. LotFrontage (18.6%) y las variables numéricas se imputaron con la mediana "
        "del fold y un indicador de ausencia. Electrical tiene un único faltante. No se eliminaron filas."
    )
    add_figure(doc, "missing_values.png", "Figura 2. Porcentaje de valores faltantes por variable.")
    doc.add_paragraph(
        "Los boxplots y scatterplots revelaron colas largas en LotArea, GrLivArea y superficies de sótano/garage. En vez de "
        "borrar observaciones posiblemente válidas, el pipeline ganador aprendio límites en los percentiles 1 y 99 dentro de "
        "cada fold. Esta decisión reduce la influencia de extremos sobre gradientes sin usar información de válidación."
    )

    doc.add_heading("1.3 Relaciones con SalePrice", level=2)
    corr_rows = [(name, f"{value:.3f}") for name, value in corr.head(10).items()]
    add_table(doc, ["Variable", "Correlación de Pearson"], corr_rows, [5200, 4160], font_size=9.5)
    doc.add_paragraph(
        f"OverallQual ({corr.iloc[0]:.3f}) y GrLivArea ({corr.iloc[1]:.3f}) dominan las relaciones lineales, seguidas por "
        "capacidad/área de garage, sótano y primer piso. Las relaciones no son enteramente lineales: la dispersión aumenta "
        "en casas costosas y Neighborhood desplaza claramente la mediana del precio. La colinealidad entre medidas de área "
        "(GarageCars/GarageArea o 1stFlrSF/TotalBsmtSF) no obliga a eliminarlas en un MLP regularizado, pero si aumenta el riesgo de sobreajuste."
    )
    add_figure(doc, "correlation_heatmap.png", "Figura 3. Mapa de correlaciones de las variables numéricas más relacionadas con SalePrice.")
    add_figure(doc, "key_scatterplots.png", "Figura 4. Relaciones de seis predictores principales con SalePrice.")

    doc.add_heading("1.4 Decisiones derivadas del EDA", level=2)
    for text in [
        "Categoricas: Missing explicito, one-hot, categorías infrecuentes agrupadas con min_frequency=5 y tolerancia a niveles no vistos.",
        "Numericas: mediana e indicador de ausencia, clipping 1/99, log1p solo para columnas no negativas con asimetría absoluta mayor a 0.75 y RobustScaler.",
        "Ingenieria de variables: superficie total, baños equivalentes, porches, antigüedad al vender, tiempo desde remodelacion e interacciones calidad-superficie.",
        "Objetivo: se compararon SalePrice original y log1p; el ganador usa escala original estandarizada durante entrenamiento.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("2. Metodología de desarrollo", level=1)
    doc.add_heading("2.1 Separacion y control de leakage", level=2)
    doc.add_paragraph(
        "Se reservó 20% como holdout interno con random_state=20260817 y estratificación por deciles de precio. El 80% restante "
        "se usó para exploración y Optuna con tres folds estratificados. Los tres candidatos más prometedores se reevaluaron "
        "en una partición nueva de cinco folds (seed 31415). Solo después de elegir el ganador se abrio el holdout. Imputadores, "
        "umbrales de asimetría, percentiles, escaladores, agrupación de categorías y one-hot se ajustaron exclusivamente en el train de cada fold."
    )
    doc.add_page_break()
    doc.add_heading("2.2 Arquitectura y entrenamiento", level=2)
    cfg_rows = [
        ("Arquitectura", "512 -> 256 -> 128 -> 64 -> 1"), ("Activacion", "GELU"),
        ("Dropout", f"{params['dropout']:.4f}"), ("Normalización", "Ninguna"),
        ("Optimizador", "Adam"), ("Learning rate", f"{params['learning_rate']:.7f}"),
        ("Weight decay", f"{params['weight_decay']:.7f}"), ("Batch size", str(params['batch_size'])),
        ("Loss", "MSE sobre SalePrice estandarizado"), ("Gradient clipping", "Norma maxima 5.0"),
        ("Scheduler", "ReduceLROnPlateau, factor 0.5"), ("Early stopping", "Paciencia 50"),
    ]
    add_table(doc, ["Componente", "Configuracion final"], cfg_rows, [3300, 6060], font_size=9.5)
    doc.add_paragraph(
        "PyTorch permitió controlar dropout, inicialización Kaiming, weight decay, clipping, scheduler y restauración del mejor checkpoint. "
        "Aunque MPS estaba disponible en el M1 Pro, CPU fue elegido por estabilidad y menor overhead en matrices pequeñas. La semilla global fue 42."
    )
    doc.add_heading("2.3 Busqueda de hiperparámetros", level=2)
    doc.add_paragraph(
        f"Optuna con TPE exploró {len(trials)} trials y podó {int((trials.state == 'PRUNED').sum())} mediante MedianPruner. Se variaron "
        "profundidad/ancho, ReLU/GELU/LeakyReLU/SiLU, dropout, normalizacion, Adam/AdamW, learning rate, weight decay, batch size, "
        "target original/log1p, Standard/RobustScaler, clipping, umbral de asimetría, frecuencia minima e ingeniería de features. "
        "La métrica objetiva siempre fue RMSE tras invertir cualquier transformación del objetivo."
    )

    doc.add_heading("3. Resultados de iteraciones", level=1)
    major_ids = ["baseline_raw", "baseline_log", "feature_engineered", "deep_low_regularization",
                 "robust_regularized", "optuna_013", "confirm_optuna_winner",
                 "confirm_robust_regularized"]
    labels = {
        "baseline_raw": "Baseline", "baseline_log": "Objetivo log1p", "feature_engineered": "Features + regularización",
        "deep_low_regularization": "Red profunda sin regularizar", "robust_regularized": "Robust + regularización",
        "optuna_013": "Ganador Optuna 3-fold", "confirm_optuna_winner": "Ganador confirmado 5-fold",
        "confirm_robust_regularized": "Alternativa robusta 5-fold",
    }
    main_results = results.set_index("experiment_id").loc[major_ids]
    rows = []
    for idx, row in main_results.iterrows():
        rows.append((labels[idx], row.architecture, row.target_transform,
                     f"{row.train_rmse:,.0f}", f"{row.cv_rmse_mean:,.0f}"))
    add_table(doc, ["Iteración / cambio", "Arquitectura", "Target", "RMSE train", "RMSE valid."],
              rows, [2700, 1900, 1200, 1780, 1780], font_size=8.2)
    cite = doc.add_paragraph("Fuente: experiments/results.csv, generado automaticamente por ejecuciones reales.")
    cite.paragraph_format.space_before = Pt(4)
    cite.paragraph_format.space_after = Pt(4)
    font_run(cite.runs[0], size=9, italic=True, color="666666")
    doc.add_paragraph(
        f"La mejora total contra el baseline fue de {baseline - confirm['cv_rmse_mean']:,.2f} USD ({improvement:.1f}%). "
        "La red profunda sin regularización alcanzó RMSE train de 8,021 pero CV de 33,963, evidencia de sobreajuste. "
        "El baseline log1p empeoró a 44,354: simetrizar el target no alineó la loss con la penalización absoluta del RMSE oficial. "
        "RobustScaler, clipping, weight decay y dropout estabilizaron el entrenamiento; Optuna encontró después una red mayor con "
        "GELU que mejoró nuevamente en los cinco folds de confirmación."
    )
    add_figure(doc, "training_baseline_raw.png", "Figura 5. Curvas representativas del MLP baseline.")
    add_figure(doc, "training_final_cv.png", "Figura 6. Curvas representativas del modelo ganador en CV.")
    doc.add_paragraph(
        f"El mejor número de épocas varió entre {min(x['best_epoch'] for x in confirm['fold_metrics'])} y "
        f"{max(x['best_epoch'] for x in confirm['fold_metrics'])}; la mediana fue {confirm['best_epoch_median']}. "
        "Esa mediana, determinada sin el holdout, se usó para reentrenar el artefacto final sobre las 1,168 filas."
    )

    doc.add_heading("4. Discusión de resultados", level=1)
    doc.add_heading("4.1 Generalización y complejidad", level=2)
    doc.add_paragraph(
        f"La CV confirmatoria fue {confirm['cv_rmse_mean']:,.2f} +/- {confirm['cv_rmse_std']:,.2f} USD y el holdout independiente "
        f"{holdout['rmse']:,.2f} USD. El holdout favorable no reemplaza la CV: contiene solo 234 casos y su composición produce una "
        "estimación más optimista. La selección se basa por tanto en la media de cinco folds. El gap train-CV del ganador "
        f"({confirm['train_rmse']:,.0f} vs. {confirm['cv_rmse_mean']:,.0f}) muestra capacidad suficiente y sobreajuste residual, "
        "pero menor inestabilidad que la red profunda sin regularización."
    )
    doc.add_heading("4.2 Análisis de errores", level=2)
    segment_rows = [(r.price_segment, int(r["count"]), f"{r.rmse:,.0f}", f"{r.mae:,.0f}", f"{r.bias:,.0f}")
                    for _, r in segments.iterrows()]
    add_table(doc, ["Segmento", "n", "RMSE", "MAE", "Sesgo real-pred."], segment_rows,
              [2300, 900, 1900, 1900, 2360], font_size=9)
    doc.add_paragraph(
        f"Los tres cuartiles inferiores tienen RMSE entre {segments.rmse.iloc[:3].min():,.0f} y {segments.rmse.iloc[:3].max():,.0f} USD, "
        f"mientras el cuartil alto sube a {segments.rmse.iloc[3]:,.0f}. El sesgo de +{segments.bias.iloc[3]:,.0f} en casas caras "
        "indica subestimación; en el segmento bajo el sesgo negativo indica sobreestimación. Este patron es coherente con regresión "
        "hacia la media y con menos ejemplos en la cola del precio."
    )
    top = largest.sort_values("absolute_error", ascending=False).head(5)
    top_rows = [(int(r.Id), f"{r.actual:,.0f}", f"{r.prediction:,.0f}", f"{r.absolute_error:,.0f}", r.Neighborhood)
                for _, r in top.iterrows()]
    add_table(doc, ["Id", "Real", "Predicho", "Error abs.", "Neighborhood"], top_rows,
              [900, 1800, 1800, 1900, 2960], font_size=9)
    add_figure(doc, "final_error_analysis.png", "Figura 7. Diagnóstico del holdout: ajuste, residuos y error por segmento.")
    doc.add_paragraph(
        f"El error medio absoluto fue {holdout['mae']:,.2f} USD y el sesgo global {holdout['bias']:,.2f} USD, cercano a cero. "
        "Los mayores errores se concentran en propiedades caras o poco comunes; no se observa un desplazamiento global fuerte, "
        "pero si heterocedasticidad: la amplitud de los residuos aumenta con la predicción."
    )
    doc.add_page_break()
    doc.add_heading("4.3 Limitaciones", level=2)
    for text in [
        "Solo hay 1,168 observaciones para unas 300 columnas procesadas; la varianza entre folds sigue siendo material.",
        "One-hot produce alta dimensionalidad y un MLP no incorpora de forma natural las particiones inductivas favorables de modelos de árbol para datos tabulares.",
        "Las colas de precio y los vecindarios raros tienen poca representación; el RMSE enfatiza precisamente esos errores grandes.",
        "La búsqueda de 17 trials es seria pero no exhaustiva; más seeds y repeated CV reducirían incertidumbre a mayor costo computacional.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("5. Conclusiones", level=1)
    doc.add_paragraph(
        f"El MLP final usa capas 512-256-128-64, GELU, dropout {params['dropout']:.3f}, Adam, learning rate "
        f"{params['learning_rate']:.7f}, weight decay {params['weight_decay']:.7f} y batch 64. Con preprocessing robusto e "
        f"ingeniería de variables alcanzó {confirm['cv_rmse_mean']:,.2f} USD de RMSE CV, una reducción de {improvement:.1f}% "
        f"contra el baseline. El holdout de {holdout['rmse']:,.2f} USD confirma que el pipeline generaliza, aunque el desempeño "
        "es peor y más variable en viviendas del cuartil superior."
    )
    doc.add_paragraph(
        "Los aprendizajes principales fueron: ajustar todo preprocessing dentro del fold, optimizar la métrica en escala original, "
        "regularizar redes profundas y confirmar los mejores trials con una partición distinta. Como mejoras futuras convendria "
        "repetir CV con más seeds, explorar embeddings categóricos, ensembles de MLPs con seeds distintas y una loss ponderada/calibracion "
        "que reduzca la subestimación de propiedades caras, siempre válidando el RMSE original."
    )

    doc.add_heading("6. Enlace al repositorio de GitHub", level=1)
    p = doc.add_paragraph("Código, artefacto final, resultados y pasos de reproducción: ")
    add_hyperlink(p, "github.com/Javiervalladares1/ia-proyecto1-mlp-house-prices",
                  "https://github.com/Javiervalladares1/ia-proyecto1-mlp-house-prices")
    doc.add_paragraph(
        "El comando de competencia es python predict.py test.csv. El modelo, preprocesador, transformador del objetivo, "
        "orden de columnas, versiones, seed y metadatos se encuentran en models/final/."
    )

    doc.add_heading("Anexo A. Entregables reproducibles", level=1)
    deliverables = [
        ("EDA", "run_eda.py; artifacts/; figures/"),
        ("Entrenamiento y búsqueda", "train.py; experiments/results.csv; experiments/optuna_trials.csv"),
        ("Modelo final", "models/final/model.pt y artefactos joblib/json"),
        ("Predicción", "predict.py test.csv"),
        ("Pruebas", "python -m unittest discover -s tests -v"),
        ("Informe", "reports/Informe_Proyecto_1_MLP.docx y .pdf"),
    ]
    add_table(doc, ["Componente", "Ubicación / comando"], deliverables, [3000, 6360], font_size=9.5)

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_report()
